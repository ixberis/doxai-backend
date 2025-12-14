
# -*- coding: utf-8 -*-
"""
backend/app/services/rag/converters/word_converter.py

Conversión de archivos Word (.docx, .doc, .odt) a texto estructurado para el pipeline RAG.

Funcionalidades:
- Convierte archivos .doc y .odt (legacy) a .docx utilizando LibreOffice (soffice).
- Extrae bloques semánticos desde el .docx con Unstructured (modo estándar o hi_res).
- Si falla, usa `python-docx` como fallback.
- Aplica limpieza OCR defensiva con `clean_ocr_text(...)` para mejorar la calidad del texto.
- Detecta y clasifica tablas en dos tipos: 'informativa' o 'formulario'.
- Reconstruye tablas densas mediante heurísticas si `split_table_rows(...)` no es suficiente.
- Detecta formularios tipo declaración: subrayados, campos tipo "<...>", cláusulas legales frecuentes.
- Preserva jerarquía documental con títulos y subtítulos.

Archivos generados:
- `.md` con texto limpio y jerárquico
- `.tables.json` con tablas estructuradas
- `.forms.json` con formularios tipo declaración

Autor: Ixchel Beristain
Última revisión: 01/08/2025
"""

import os
import re
from typing import Optional, List, Dict, Union, Callable
from docx import Document
from unstructured.documents.elements import Title

from .doc_legacy_converter import convert_doc_to_docx
# TODO: Implementar módulo RAG completo
# from app.modules.rag.utils.unstructured_parser_service import parse_with_unstructured
# from app.modules.rag.utils.table_reconstruction_service import reconstruct_table_from_raw
# from app.modules.rag.utils.text_structure_service import enhanced_elements_to_text_with_titles
# from app.modules.rag.utils.ocr_cleaning_service import clean_ocr_text


def extract_text_with_docx(path: str) -> str:
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                tables.append(" | ".join(cells))
    return "\n".join(paragraphs + tables)


def classify_table(table: List[List[str]]) -> str:
    keywords = [
        "precio", "cantidad", "firma", "total", "nombre", "unidad",
        "cargo", "producto", "%", "ods", "pnd", "recomendacion",
        "objetivo", "avance", "entregable", "hito", "marca"
    ]
    header_row = table[0] if table else []
    text = " ".join(header_row).lower()
    return "formulario" if any(k in text for k in keywords) else "informativa"


def detect_form_declaration(text: str) -> List[Dict[str, Union[str, List[str]]]]:
    forms = []
    lines = text.splitlines()

    for line in lines:
        if any(p in line.lower() for p in [
            "nombre del prestador", "nombre del representante legal",
            "lugar y fecha", "firma del"
        ]):
            forms.append({
                "type": "formulario_declarativo",
                "raw": line.strip(),
                "fields": [w.strip() for w in re.findall(r"[A-Z]{3,}(?: [A-Z]{3,})*", line)]
            })

    pattern = re.compile(r"(.*?__(?:_+).*?)", re.DOTALL)
    placeholder = re.compile(r"<[^>]+>")
    declarative_clause = re.compile(r"bajo protesta de decir verdad", re.IGNORECASE)

    for match in pattern.finditer(text):
        block = match.group(1).strip()
        if len(block.split()) > 4:
            forms.append({
                "type": "formulario_declarativo",
                "raw": block,
                "fields": re.findall(r"(?:__+|<[^>]+>)", block)
            })

    for match in placeholder.finditer(text):
        forms.append({
            "type": "formulario_declarativo",
            "raw": match.group(0),
            "fields": [match.group(0)]
        })

    if declarative_clause.search(text):
        forms.append({
            "type": "formulario_declarativo",
            "raw": text.strip(),
            "fields": ["bajo protesta de decir verdad"]
        })

    return forms


def split_table_rows(raw_text: str) -> List[List[str]]:
    rows = []
    blocks = raw_text.splitlines()
    for line in blocks:
        line = line.strip()
        if not line:
            continue

        if "|" in line:
            cols = [c.strip() for c in line.split("|")]
        elif "\t" in line:
            cols = [c.strip() for c in line.split("\t")]
        elif re.search(r"\s{3,}", line):
            cols = [c.strip() for c in re.split(r"\s{3,}", line)]
        else:
            cols = line.split()

        if len(cols) >= 2:
            rows.append(cols)

    return rows


def convert_word_to_text(
    file_path: str, 
    progress_cb: Optional[Callable[[str, float, Optional[Dict]], None]] = None,
    job_id: Optional[str] = None
) -> Optional[Dict[str, Union[str, List]]]:
    """
    Convierte archivos Word (.docx, .doc, .odt) a texto estructurado para el pipeline RAG.
    
    Args:
        file_path: Ruta del archivo Word
        progress_cb: Callback opcional para reporte de progreso (stage, percent, meta)
        job_id: ID opcional del job para logging contextual
    
    Returns:
        Dict con 'text', 'tables', 'forms' o None si falla la conversión
    """
    if progress_cb:
        progress_cb("preparing", 0.05, {"message": "Iniciando conversión Word"})
    
    ext = os.path.splitext(file_path)[1].lower()

    if ext in [".doc", ".odt"]:
        if progress_cb:
            progress_cb("converting_legacy", 0.15, {"message": f"Convirtiendo formato legacy {ext} a DOCX"})
        print(f"📄 Convirtiendo archivo legacy ({ext}) a .docx: {file_path}")
        converted_path = convert_doc_to_docx(file_path)
        print(f"📁 Ruta convertida: {converted_path}")
        if not converted_path:
            print("❌ Error al convertir a .docx")
            if progress_cb:
                progress_cb("error", 0.0, {"error": "Conversión legacy falló"})
            return None
        file_path = converted_path

    def process_elements(elements):
        if progress_cb:
            progress_cb("processing", 0.6, {"message": "Procesando elementos extraídos", "elements_count": len(elements)})
        
        text = enhanced_elements_to_text_with_titles(elements)
        tables = []
        forms = []

        for e in elements:
            raw_text = clean_ocr_text(e.text.strip())

            if e.category == "Table" or (
                e.category == "Text" and len(raw_text.splitlines()) > 2 and (":" in raw_text or "%" in raw_text)
            ):
                structured_rows = split_table_rows(raw_text)
                if structured_rows:
                    tables.append({
                        "rows": structured_rows,
                        "type": classify_table(structured_rows)
                    })
                else:
                    reconstructed = reconstruct_table_from_raw(raw_text)
                    tables.append(reconstructed)

            if e.category == "Text":
                detected = detect_form_declaration(raw_text)
                if detected:
                    forms.extend(detected)

        return text, tables, forms

    if progress_cb:
        progress_cb("extracting", 0.25, {"message": "Extrayendo contenido con Unstructured"})
        
    print(f"🔍 Intentando con Unstructured (modo estándar): {file_path}")
    elements = parse_with_unstructured(file_path)
    if elements:
        text, tables, forms = process_elements(elements)
        print(f"✅ Bloques detectados: {len(elements)} (modo estándar)")
        
        if progress_cb:
            progress_cb("completed", 0.95, {
                "message": "Conversión Word completada",
                "text_length": len(text),
                "tables_count": len(tables),
                "forms_count": len(forms)
            })
        
        return {
            "text": text,
            "tables": tables,
            "forms": forms
        }

    if progress_cb:
        progress_cb("fallback_hires", 0.4, {"message": "Reintentando con estrategia hi-res"})
        
    print("⚠️ Modo estándar falló, reintentando con strategy='hi_res'")
    from unstructured.partition.auto import partition
    try:
        elements_hi_res = partition(file_path, strategy="hi_res", infer_table_structure=True)
        if elements_hi_res:
            text, tables, forms = process_elements(elements_hi_res)
            print(f"✅ Bloques detectados: {len(elements_hi_res)} (hi_res)")
            
            if progress_cb:
                progress_cb("completed", 0.95, {
                    "message": "Conversión Word completada (hi-res)",
                    "text_length": len(text),
                    "tables_count": len(tables),
                    "forms_count": len(forms)
                })
            
            return {
                "text": text,
                "tables": tables,
                "forms": forms
            }
    except Exception as e:
        print(f"❌ Error en Unstructured hi_res: {e}")
        if progress_cb:
            progress_cb("error", 0.0, {"error": f"Hi-res falló: {str(e)}"})

    if progress_cb:
        progress_cb("fallback_docx", 0.5, {"message": "Fallback con python-docx"})
        
    print("⚠️ Extracción con Unstructured fallida. Recurriendo a python-docx.")
    try:
        text_fallback = extract_text_with_docx(file_path)
        if text_fallback.strip():
            forms = detect_form_declaration(clean_ocr_text(text_fallback))
            print(f"✅ Texto extraído con fallback (python-docx). Longitud: {len(text_fallback)}")
            
            if progress_cb:
                progress_cb("completed", 0.95, {
                    "message": "Conversión completada con python-docx",
                    "text_length": len(text_fallback),
                    "tables_count": 0,
                    "forms_count": len(forms)
                })
            
            return {
                "text": text_fallback,
                "tables": [],
                "forms": forms
            }
    except Exception as e:
        print(f"❌ Fallback con python-docx falló: {e}")
        if progress_cb:
            progress_cb("error", 0.0, {"error": f"Fallback python-docx falló: {str(e)}"})

    print("❌ Todos los métodos de extracción fallaron.")
    if progress_cb:
        progress_cb("error", 0.0, {"error": "Todos los métodos de conversión fallaron"})
    return None
# Fin del archivo word_converter.py







